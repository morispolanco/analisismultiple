import streamlit as st
import requests
import docx
from datetime import datetime
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import matplotlib.pyplot as plt
from wordcloud import WordCloud
import io
import os
import json

# Especificar un directorio personalizado para los datos de NLTK
nltk_data_path = os.path.join(os.getcwd(), "nltk_data")
if not os.path.exists(nltk_data_path):
    os.makedirs(nltk_data_path)
nltk.data.path.append(nltk_data_path)

# Descargar recursos de NLTK solo si no están presentes
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', download_dir=nltk_data_path)
try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt_tab', download_dir=nltk_data_path)
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', download_dir=nltk_data_path)

# Configuración de la página
st.set_page_config(page_title="Análisis Literario Multidisciplinario", layout="wide")

# Título y descripción
st.title("Análisis Literario Multidisciplinario")
st.write("Esta aplicación combina estudios literarios, históricos y antropológicos para analizar la obra de un autor.")

# Sidebar para inputs
with st.sidebar:
    st.header("Configuración")
    author_name = st.text_input("Nombre del autor", "Gabriel García Márquez")
    work_title = st.text_input("Obra principal", "Cien años de soledad")
    analyze_button = st.button("Analizar")

# Función para llamar a la API de Gemini
def get_essay_section_from_api(author, work, section_prompt):
    url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
    api_key = st.secrets['GEMINI_API_KEY']
    full_url = f"{url}?key={api_key}"
    
    headers = {
        "Content-Type": "application/json"
    }
    
    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [
                    {
                        "text": section_prompt
                    }
                ]
            }
        ],
        "generationConfig": {
            "temperature": 1,
            "topK": 40,
            "topP": 0.95,
            "maxOutputTokens": 8192,  # Máximo permitido por Gemini 2.0 Flash
            "responseMimeType": "text/plain"
        }
    }
    
    response = requests.post(full_url, headers=headers, json=payload)
    if response.status_code == 200:
        return response.json()['candidates'][0]['content']['parts'][0]['text']
    else:
        return f"Error en la API: {response.status_code} - {response.text}"

# Función para generar el ensayo completo
def get_full_essay(author, work):
    # Dividir el ensayo en secciones para superar las 6000 palabras
    prompts = [
        f"Escribe una introducción extensa (mínimo 2000 palabras) para un ensayo académico que combine estudios literarios, históricos y antropológicos sobre la obra '{work}' de {author}. Incluye contexto y citas en formato APA.",
        f"Escribe un análisis detallado (mínimo 3000 palabras) que combine estudios literarios, históricos y antropológicos sobre la obra '{work}' de {author}. Profundiza en temas, simbolismos y contexto, con citas en formato APA.",
        f"Escribe unas conclusiones extensas (mínimo 1000 palabras) para un ensayo académico sobre la obra '{work}' de {author}, resumiendo el análisis literario, histórico y antropológico, con citas en formato APA."
    ]
    
    essay_parts = []
    for prompt in prompts:
        section = get_essay_section_from_api(author, work, prompt)
        if "Error en la API" in section:
            return section  # Retorna el error si ocurre
        essay_parts.append(section)
    
    full_essay = "\n\n".join(essay_parts)
    
    # Verificar longitud
    word_count = len(full_essay.split())
    if word_count < 6000:
        st.warning(f"El ensayo tiene {word_count} palabras, menos de las 6000 solicitadas debido a límites de la API.")
    
    return full_essay

# Función para análisis lingüístico
def linguistic_analysis(text):
    tokens = word_tokenize(text.lower())
    stop_words = set(stopwords.words('spanish'))
    filtered_tokens = [word for word in tokens if word.isalnum() and word not in stop_words]
    
    # Frecuencia de palabras
    freq = nltk.FreqDist(filtered_tokens)
    return freq, filtered_tokens

# Función para crear documento Word
def create_word_doc(essay, author, work, freq):
    doc = docx.Document()
    doc.add_heading(f"Análisis de {work} de {author}", 0)
    
    # Introducción
    doc.add_heading("Introducción", level=1)
    doc.add_paragraph(essay[:10000])  # Ajustado para secciones más largas
    
    # Análisis lingüístico
    doc.add_heading("Análisis Lingüístico", level=1)
    doc.add_paragraph("Resultados del análisis computacional del texto:")
    for word, count in freq.most_common(10):
        doc.add_paragraph(f"{word}: {count} ocurrencias")
    
    # Cuerpo del ensayo
    doc.add_heading("Análisis Multidisciplinario", level=1)
    doc.add_paragraph(essay[10000:-5000])
    
    # Conclusiones
    doc.add_heading("Conclusiones", level=1)
    doc.add_paragraph(essay[-5000:])
    
    # Referencias
    doc.add_heading("Referencias", level=1)
    doc.add_paragraph("García, J. (2020). Contexto histórico de la literatura latinoamericana. Journal of Latin American Studies, 45(3), 234-256.")
    
    return doc

# Lógica principal
if analyze_button:
    with st.spinner("Generando análisis (esto puede tomar varios minutos)..."):
        # Obtener ensayo completo
        essay = get_full_essay(author_name, work_title)
        
        if "Error en la API" in essay:
            st.error(essay)
        else:
            # Análisis lingüístico
            freq, tokens = linguistic_analysis(essay)
            
            # Visualizaciones
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("Nube de palabras")
                wordcloud = WordCloud(width=800, height=400).generate_from_frequencies(dict(freq))
                plt.figure(figsize=(10, 5))
                plt.imshow(wordcloud, interpolation='bilinear')
                plt.axis('off')
                st.pyplot(plt)
            
            with col2:
                st.subheader("Palabras más frecuentes")
                plt.figure(figsize=(10, 5))
                freq.plot(20, cumulative=False)
                st.pyplot(plt)
            
            # Mostrar ensayo
            st.subheader("Ensayo Generado")
            st.write(essay[:2000] + "... [continúa]")
            
            # Exportar a Word
            doc = create_word_doc(essay, author_name, work_title, freq)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.download_button(
                label="Descargar ensayo en Word",
                data=buffer,
                file_name=f"Analisis_{author_name}_{work_title}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Footer
st.markdown("---")
st.write("Desarrollado con Streamlit y Gemini API - Marzo 2025")
